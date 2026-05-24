import os
import sqlite3
from datetime import date, datetime
from pathlib import Path

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

APP_DIR = Path(__file__).parent
DATA_DIR = APP_DIR / "data"
EXPORT_DIR = APP_DIR / "exports"
DB_PATH = DATA_DIR / "records.db"
DATA_DIR.mkdir(exist_ok=True)
EXPORT_DIR.mkdir(exist_ok=True)

DEFAULT_COURSES = [
    ("20100-1008", "งานนิวเมติกส์และไฮดรอลิกส์เบื้องต้น"),
]
DEFAULT_CLASSES = ["ชฟ.3/1", "ชฟ.3/2", "ชฟ.2/1", "ชฟ.2/2"]
DEFAULT_TEACHER = "นายบุรินทร์ พุฒทอง"


def connect():
    return sqlite3.connect(DB_PATH)


def init_db():
    conn = connect()
    c = conn.cursor()
    c.execute("""
        CREATE TABLE IF NOT EXISTS records (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            created_at TEXT,
            teach_date TEXT,
            week_no INTEGER,
            lesson_no INTEGER,
            course TEXT,
            class_group TEXT,
            teacher TEXT,
            student_count INTEGER,
            topic TEXT,
            activity TEXT,
            problem TEXT,
            solution TEXT,
            result TEXT,
            class_line TEXT
        )
    """)
    # เพิ่มคอลัมน์สำหรับบรรทัดชั้น/ห้อง (สำหรับฐานข้อมูลเวอร์ชันเก่า)
    cols = [r[1] for r in c.execute("PRAGMA table_info(records)").fetchall()]
    if "class_line" not in cols:
        c.execute("ALTER TABLE records ADD COLUMN class_line TEXT DEFAULT ''")

    c.execute("""
        CREATE TABLE IF NOT EXISTS courses (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            course_code TEXT,
            course_name TEXT,
            active INTEGER DEFAULT 1,
            UNIQUE(course_code, course_name)
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS classes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            class_name TEXT UNIQUE,
            active INTEGER DEFAULT 1
        )
    """)
    c.execute("""
        CREATE TABLE IF NOT EXISTS settings (
            key TEXT PRIMARY KEY,
            value TEXT
        )
    """)

    for code, name in DEFAULT_COURSES:
        c.execute("INSERT OR IGNORE INTO courses(course_code, course_name, active) VALUES (?, ?, 1)", (code, name))
    for class_name in DEFAULT_CLASSES:
        c.execute("INSERT OR IGNORE INTO classes(class_name, active) VALUES (?, 1)", (class_name,))
    c.execute("INSERT OR IGNORE INTO settings(key, value) VALUES ('teacher', ?)", (DEFAULT_TEACHER,))
    conn.commit()
    conn.close()


def get_teacher():
    conn = connect()
    row = conn.execute("SELECT value FROM settings WHERE key='teacher'").fetchone()
    conn.close()
    return row[0] if row else DEFAULT_TEACHER


def set_teacher(name):
    conn = connect()
    conn.execute("INSERT OR REPLACE INTO settings(key, value) VALUES ('teacher', ?)", (name,))
    conn.commit()
    conn.close()


def load_courses(active_only=True):
    conn = connect()
    where = "WHERE active=1" if active_only else ""
    df = pd.read_sql_query(f"SELECT * FROM courses {where} ORDER BY id DESC", conn)
    conn.close()
    return df


def load_classes(active_only=True):
    conn = connect()
    where = "WHERE active=1" if active_only else ""
    df = pd.read_sql_query(f"SELECT * FROM classes {where} ORDER BY id DESC", conn)
    conn.close()
    return df


def add_course(code, name):
    conn = connect()
    conn.execute("INSERT OR IGNORE INTO courses(course_code, course_name, active) VALUES (?, ?, 1)", (code.strip(), name.strip()))
    conn.commit()
    conn.close()


def update_course_active(course_id, active):
    conn = connect()
    conn.execute("UPDATE courses SET active=? WHERE id=?", (1 if active else 0, course_id))
    conn.commit()
    conn.close()


def add_class(class_name):
    conn = connect()
    conn.execute("INSERT OR IGNORE INTO classes(class_name, active) VALUES (?, 1)", (class_name.strip(),))
    conn.commit()
    conn.close()


def update_class_active(class_id, active):
    conn = connect()
    conn.execute("UPDATE classes SET active=? WHERE id=?", (1 if active else 0, class_id))
    conn.commit()
    conn.close()


def save_record(data):
    conn = connect()
    c = conn.cursor()
    c.execute("""
        INSERT INTO records (
            created_at, teach_date, week_no, lesson_no, course, class_group, teacher,
            student_count, topic, activity, problem, solution, result, class_line
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        data["teach_date"].isoformat(), data["week_no"], data["lesson_no"], data["course"],
        data["class_group"], data["teacher"], data["student_count"], data["topic"],
        data["activity"], data["problem"], data["solution"], data["result"], data.get("class_line", ""),
    ))
    conn.commit()
    record_id = c.lastrowid
    conn.close()
    return record_id


def load_records():
    conn = connect()
    df = pd.read_sql_query("SELECT * FROM records ORDER BY teach_date DESC, id DESC", conn)
    conn.close()
    return df


def thai_date(d):
    if isinstance(d, str):
        d = datetime.fromisoformat(d).date()
    months = ["", "มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน", "พฤษภาคม", "มิถุนายน", "กรกฎาคม", "สิงหาคม", "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม"]
    return f"{d.day} {months[d.month]} {d.year + 543}"


def set_cell_shading(cell, fill="EAF2F8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=14):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "TH Sarabun New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_section(doc, title, body):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.name = "TH Sarabun New"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
    r.font.size = Pt(16)
    body_p = doc.add_paragraph()
    r2 = body_p.add_run(str(body).strip() if str(body).strip() else "-")
    r2.font.name = "TH Sarabun New"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
    r2.font.size = Pt(15)


def split_course(course):
    if " | " in course:
        return (course.split(" | ", 1) + [""])[:2]
    return "", course



def set_docx_font(run, size=16, bold=False):
    run.font.name = "TH Sarabun New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
    run.font.size = Pt(size)
    run.bold = bold


def add_center_line(doc, text, size=16, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    set_docx_font(r, size=size, bold=bold)
    return p


def add_body_text(doc, text, size=16, bold=False, space_after=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.0
    for idx, line in enumerate(str(text).split("\n")):
        if idx:
            p.add_run().add_break()
        r = p.add_run(line)
        set_docx_font(r, size=size, bold=bold)
    return p


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def export_docx(data, output_path):
    """สร้าง Word ตามฟอร์มต้นฉบับ: ไม่มีตาราง ไม่มีโลโก้ ไม่มีลายเซ็น"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.35)
    section.right_margin = Cm(2.1)

    style = doc.styles['Normal']
    style.font.name = 'TH Sarabun New'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
    style.font.size = Pt(16)

    course_code, course_name = split_course(data["course"])
    class_line = data.get("class_line") or f"ช่างไฟฟ้ากำลัง {data['class_group']}"

    add_center_line(doc, "บันทึกหลังการจัดการเรียนรู้", size=16, bold=False)
    add_center_line(doc, f"รหัสวิชา {course_code} รายวิชา {course_name}", size=16)
    add_center_line(doc, class_line, size=16)
    add_center_line(doc, f"ครูผู้สอน {data['teacher']} จำนวน", size=16)
    add_center_line(doc, f"วันที่ {thai_date(data['teach_date'])} สัปดาห์ที่ {data['week_no']} จำนวน {data['student_count']} คน", size=16)

    doc.add_paragraph("")
    add_body_text(doc, "หัวข้อเรื่อง/เนื้อหาสาระ/การอบรม/ให้คำปรึกษา/บันทึกการสอน :", size=16)
    add_body_text(doc, str(data["topic"]).strip() or "-", size=16)

    for title, body in [
        ("รายละเอียด/กิจกรรม", data["activity"]),
        ("ปัญหา/อุปสรรค", data["problem"]),
        ("แนวทางการแก้ไขและปรับปรุง", data["solution"]),
        ("ผลการจัดการเรียนรู้", data["result"]),
    ]:
        add_horizontal_line(doc)
        add_body_text(doc, title, size=16)
        add_body_text(doc, str(body).strip() or "-", size=16)

    doc.save(output_path)


def find_thai_font():
    candidates = [
        "C:/Windows/Fonts/THSarabunNew.ttf",
        "C:/Windows/Fonts/THSarabun.ttf",
        "C:/Windows/Fonts/tahoma.ttf",
        "/usr/share/fonts/truetype/noto/NotoSansThai-Regular.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for path in candidates:
        if os.path.exists(path):
            return path
    return None


def export_pdf(data, output_path):
    """สร้าง PDF ให้หน้าตาใกล้ฟอร์ม Word เดิม ไม่มีโลโก้ ไม่มีลายเซ็น"""
    font_path = find_thai_font()
    font_name = "Helvetica"
    if font_path:
        try:
            pdfmetrics.registerFont(TTFont("ThaiFont", font_path))
            font_name = "ThaiFont"
        except Exception:
            pass

    doc = SimpleDocTemplate(
        str(output_path), pagesize=A4,
        rightMargin=2.1*cm, leftMargin=2.35*cm,
        topMargin=1.55*cm, bottomMargin=1.5*cm
    )
    styles = getSampleStyleSheet()
    normal = ParagraphStyle("ThaiNormal", parent=styles["Normal"], fontName=font_name, fontSize=13.5, leading=20, spaceAfter=0)
    center = ParagraphStyle("ThaiCenter", parent=normal, alignment=1, fontSize=13.5, leading=19)
    section_style = ParagraphStyle("Section", parent=normal, fontSize=13.5, leading=20, spaceBefore=2)

    course_code, course_name = split_course(data["course"])
    class_line = data.get("class_line") or f"ช่างไฟฟ้ากำลัง {data['class_group']}"

    def esc(x):
        return str(x).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")

    story = []
    for line in [
        "บันทึกหลังการจัดการเรียนรู้",
        f"รหัสวิชา {course_code} รายวิชา {course_name}",
        class_line,
        f"ครูผู้สอน {data['teacher']} จำนวน",
        f"วันที่ {thai_date(data['teach_date'])} สัปดาห์ที่ {data['week_no']} จำนวน {data['student_count']} คน",
    ]:
        story.append(Paragraph(esc(line), center))
    story.append(Spacer(1, 0.65*cm))

    story.append(Paragraph("หัวข้อเรื่อง/เนื้อหาสาระ/การอบรม/ให้คำปรึกษา/บันทึกการสอน :", section_style))
    story.append(Paragraph(esc(str(data["topic"]).strip() or "-"), normal))

    for title, body in [
        ("รายละเอียด/กิจกรรม", data["activity"]),
        ("ปัญหา/อุปสรรค", data["problem"]),
        ("แนวทางการแก้ไขและปรับปรุง", data["solution"]),
        ("ผลการจัดการเรียนรู้", data["result"]),
    ]:
        story.append(Spacer(1, 0.32*cm))
        line = Table([[""]], colWidths=[16.55*cm], rowHeights=[0.01*cm])
        line.setStyle(TableStyle([('LINEABOVE', (0,0), (-1,-1), 0.5, colors.HexColor('#999999'))]))
        story.append(line)
        story.append(Paragraph(title, section_style))
        story.append(Paragraph(esc(str(body).strip() or "-"), normal))

    doc.build(story)


def build_data_from_row(row):
    return dict(
        teach_date=datetime.fromisoformat(row["teach_date"]).date(),
        week_no=int(row["week_no"]),
        lesson_no=int(row["lesson_no"]),
        course=row["course"],
        class_group=row["class_group"],
        teacher=row["teacher"],
        student_count=int(row["student_count"]),
        topic=row["topic"], activity=row["activity"], problem=row["problem"], solution=row["solution"], result=row["result"],
        class_line=row.get("class_line", "") if hasattr(row, "get") else "",
    )


def make_export_files(data, export_type):
    safe_date = data["teach_date"].strftime("%Y%m%d")
    base = f"บันทึกหลังสอน_{data['class_group']}_ครั้งที่{data['lesson_no']}_{safe_date}".replace("/", "-")
    paths = []
    if export_type in ["Word", "Word และ PDF"]:
        docx_path = EXPORT_DIR / f"{base}.docx"
        export_docx(data, docx_path)
        paths.append(("Word", docx_path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))
    if export_type in ["PDF", "Word และ PDF"]:
        pdf_path = EXPORT_DIR / f"{base}.pdf"
        export_pdf(data, pdf_path)
        paths.append(("PDF", pdf_path, "application/pdf"))
    return paths


init_db()
st.set_page_config(page_title="บันทึกหลังการสอน", page_icon="📘", layout="wide")
st.title("📘 ระบบบันทึกหลังการจัดการเรียนรู้")
st.caption("ใช้ฟอร์ม Word แบบต้นฉบับ เลือกวันที่ Export เป็น Word/PDF และเก็บประวัติย้อนหลัง")

menu = st.sidebar.radio("เมนู", ["กรอกบันทึกหลังสอน", "ประวัติย้อนหลัง", "จัดการรายวิชา/ห้องเรียน"])

if menu == "จัดการรายวิชา/ห้องเรียน":
    st.subheader("⚙️ จัดการข้อมูลพื้นฐาน")
    teacher = st.text_input("ชื่อครูผู้สอนเริ่มต้น", get_teacher())
    if st.button("บันทึกชื่อครู"):
        set_teacher(teacher)
        st.success("บันทึกชื่อครูแล้ว")

    st.divider()
    st.markdown("### เพิ่มรายวิชา")
    col1, col2, col3 = st.columns([1.5, 3, 1])
    with col1:
        new_code = st.text_input("รหัสวิชา", placeholder="เช่น 20100-1008")
    with col2:
        new_name = st.text_input("ชื่อรายวิชา", placeholder="เช่น งานนิวเมติกส์และไฮดรอลิกส์เบื้องต้น")
    with col3:
        st.write("")
        st.write("")
        if st.button("➕ เพิ่มรายวิชา"):
            if new_code.strip() and new_name.strip():
                add_course(new_code, new_name)
                st.success("เพิ่มรายวิชาแล้ว")
                st.rerun()
            else:
                st.warning("กรอกรหัสวิชาและชื่อรายวิชาก่อน")

    courses_df = load_courses(active_only=False)
    if not courses_df.empty:
        st.markdown("#### รายวิชาทั้งหมด")
        for _, row in courses_df.iterrows():
            c1, c2, c3 = st.columns([1.5, 4, 1])
            c1.write(row["course_code"])
            c2.write(row["course_name"])
            label = "ปิดใช้" if int(row["active"]) == 1 else "เปิดใช้"
            if c3.button(label, key=f"course_{row['id']}"):
                update_course_active(int(row["id"]), not bool(row["active"]))
                st.rerun()

    st.divider()
    st.markdown("### เพิ่มห้อง/กลุ่มเรียน")
    col1, col2 = st.columns([3, 1])
    with col1:
        new_class = st.text_input("ชื่อห้อง/กลุ่มเรียน", placeholder="เช่น ชฟ.3/1")
    with col2:
        st.write("")
        st.write("")
        if st.button("➕ เพิ่มห้อง"):
            if new_class.strip():
                add_class(new_class)
                st.success("เพิ่มห้องแล้ว")
                st.rerun()
            else:
                st.warning("กรอกชื่อห้องก่อน")

    classes_df = load_classes(active_only=False)
    if not classes_df.empty:
        st.markdown("#### ห้อง/กลุ่มเรียนทั้งหมด")
        for _, row in classes_df.iterrows():
            c1, c2 = st.columns([5, 1])
            c1.write(row["class_name"])
            label = "ปิดใช้" if int(row["active"]) == 1 else "เปิดใช้"
            if c2.button(label, key=f"class_{row['id']}"):
                update_class_active(int(row["id"]), not bool(row["active"]))
                st.rerun()

elif menu == "กรอกบันทึกหลังสอน":
    courses_df = load_courses(active_only=True)
    classes_df = load_classes(active_only=True)
    course_options = [f"{r.course_code} | {r.course_name}" for r in courses_df.itertuples()]
    class_options = classes_df["class_name"].tolist()

    if not course_options or not class_options:
        st.warning("กรุณาเพิ่มรายวิชาและห้องเรียนก่อน")
    else:
        with st.form("record_form"):
            col1, col2, col3 = st.columns(3)
            with col1:
                teach_date = st.date_input("วันที่สอน", value=date.today())
                week_no = st.number_input("สัปดาห์ที่", min_value=1, max_value=30, value=1)
            with col2:
                lesson_no = st.selectbox("การเรียนครั้งที่", list(range(1, 31)))
                student_count = st.number_input("จำนวนผู้เรียน", min_value=0, max_value=100, value=18)
            with col3:
                course = st.selectbox("รายวิชา", course_options)
                class_group = st.selectbox("ห้อง/กลุ่มเรียน", class_options)

            default_class_line = f"ช่างไฟฟ้ากำลัง ไฟฟ้ากำลัง/1 2565 ({class_group} )"
            class_line = st.text_input("บรรทัดชั้น/กลุ่มเรียน (ตามฟอร์ม Word)", default_class_line)
            teacher = st.text_input("ครูผู้สอน", get_teacher())
            export_type = st.radio("ต้องการดาวน์โหลดไฟล์แบบไหน", ["Word และ PDF", "Word", "PDF"], horizontal=True)

            topic = st.text_area("หัวข้อเรื่อง/เนื้อหาสาระ/การอบรม/ให้คำปรึกษา/บันทึกการสอน", height=150)
            activity = st.text_area("รายละเอียด/กิจกรรม", height=100)
            problem = st.text_area("ปัญหา/อุปสรรค", height=100)
            solution = st.text_area("แนวทางการแก้ไขและปรับปรุง", height=100)
            result = st.text_area("ผลการจัดการเรียนรู้", height=100)
            submitted = st.form_submit_button("💾 บันทึกและสร้างไฟล์")

        if submitted:
            data = dict(
                teach_date=teach_date, week_no=int(week_no), lesson_no=int(lesson_no), course=course,
                class_group=class_group, class_line=class_line, teacher=teacher, student_count=int(student_count),
                topic=topic, activity=activity, problem=problem, solution=solution, result=result,
            )
            record_id = save_record(data)
            st.success(f"บันทึกสำเร็จ เลขที่รายการ {record_id}")
            for label, path, mime in make_export_files(data, export_type):
                with open(path, "rb") as f:
                    st.download_button(f"⬇️ ดาวน์โหลด {label}", f, file_name=path.name, mime=mime)

elif menu == "ประวัติย้อนหลัง":
    st.subheader("📚 ประวัติย้อนหลัง")
    df = load_records()
    if df.empty:
        st.info("ยังไม่มีข้อมูล")
    else:
        col1, col2, col3 = st.columns(3)
        with col1:
            class_filter = st.selectbox("กรองตามห้อง", ["ทั้งหมด"] + sorted(df["class_group"].dropna().unique().tolist()))
        with col2:
            course_filter = st.selectbox("กรองตามรายวิชา", ["ทั้งหมด"] + sorted(df["course"].dropna().unique().tolist()))
        with col3:
            export_type_history = st.selectbox("สร้างไฟล์ย้อนหลัง", ["Word และ PDF", "Word", "PDF"])

        show = df.copy()
        if class_filter != "ทั้งหมด":
            show = show[show["class_group"] == class_filter]
        if course_filter != "ทั้งหมด":
            show = show[show["course"] == course_filter]

        st.dataframe(show[["id", "teach_date", "week_no", "lesson_no", "course", "class_group", "student_count"]], use_container_width=True)
        st.download_button("⬇️ ดาวน์โหลดประวัติเป็น CSV", show.to_csv(index=False).encode("utf-8-sig"), file_name="post_teaching_records.csv", mime="text/csv")

        st.markdown("### Export รายการย้อนหลัง")
        ids = show["id"].tolist()
        selected_id = st.selectbox("เลือกรายการจากเลข ID", ids)
        if st.button("สร้างไฟล์จากประวัติ"):
            row = show[show["id"] == selected_id].iloc[0].to_dict()
            data = build_data_from_row(row)
            for label, path, mime in make_export_files(data, export_type_history):
                with open(path, "rb") as f:
                    st.download_button(f"⬇️ ดาวน์โหลด {label}", f, file_name=path.name, mime=mime, key=f"history_{label}_{selected_id}")
